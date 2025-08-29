import os
import json
import time
from django.http import StreamingHttpResponse
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_core.callbacks import BaseCallbackHandler
from api_chat_bot import settings
from langchain_community.docstore.document import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from typing import List, Tuple
import re
import docx
from .models import create_document_embedding

class StreamingCallbackHandler(BaseCallbackHandler):
    """Custom callback handler for streaming responses."""
    
    def __init__(self):
        self.response_tokens = []
    
    def on_llm_new_token(self, token: str, **kwargs):
        """Called when a new token is generated."""
        self.response_tokens.append(token)
    
    def get_streaming_response(self):
        """Yield tokens as they are generated."""
        for token in self.response_tokens:
            yield f"data: {json.dumps({'content': token, 'type': 'token'})}\n\n"
        
        # Send end signal
        yield f"data: {json.dumps({'type': 'end'})}\n\n"


class ChatService:
    def __init__(self):
        api_key = settings.OPENAI_API_KEY
        if not api_key:
            raise ValueError("OPENAI_API_KEY environment variable is not set")
        
        model = "gpt-4o-mini"

        self.llm = ChatOpenAI(
            model=model,
            temperature=0.7,
            streaming=True,
            openai_api_key=api_key,
            verbose=False,
        )
        
    def chat(self, input_data):
        return self.stream_chat(input_data)

    def stream_chat(self, input_data):
        """
        Stream chat response using LangChain and OpenAI with whitepaper context.
        
        Args:
            input_data: Dictionary containing user input, typically {'message': 'user message'}
        """
        user_message = input_data.get('message', '')
        
        if not user_message:
            def error_stream():
                yield f"data: {json.dumps({'error': 'No message provided', 'type': 'error'})}\n\n"
            
            response = StreamingHttpResponse(error_stream(), content_type='text/event-stream')
            response['Cache-Control'] = 'no-cache'
            return response

        def event_stream():
            try:
                # Get relevant context from vector database
                relevant_docs = self.search_whitepaper_embeddings(user_message, top_k=3)
                
                # Combine relevant context with fallback content
                context_content = ""
                if relevant_docs:
                    context_content = "\n\n".join([doc.page_content for doc in relevant_docs])
                else:
                    # Fallback to the full whitepaper content if no relevant chunks found
                    context_content = ""
                
                # Create system message with whitepaper context
                system_message = f"""You are a helpful AI assistant with access to a whitepaper about Tosi Growth Holding. 
                Use the following relevant whitepaper content to answer questions accurately and comprehensively:

                {context_content}

                When answering questions:
                1. Base your responses on the whitepaper content provided above
                2. If the question is not related to the whitepaper, politely redirect to whitepaper-related topics
                3. Provide specific information from the whitepaper when possible
                4. Be helpful and informative while staying within the context of the whitepaper
                5. If you don't have enough information to answer the question, say so and suggest what information might be available
                """

                print("=system_message===============================================")
                print(system_message)
                print("================================================")
                
                # Create messages with system context
                messages = [
                    SystemMessage(content=system_message),
                    HumanMessage(content=user_message)
                ]
                
                # Stream the response
                for chunk in self.llm.stream(messages):
                    if hasattr(chunk, 'content') and chunk.content:
                        yield f"data: {json.dumps({'content': chunk.content, 'type': 'token'})}\n\n"
                    time.sleep(0.01)
                
                # Send end signal
                yield f"data: {json.dumps({'type': 'end'})}\n\n"
                
            except Exception as e:
                # Handle specific OpenAI errors
                error_message = str(e)
                
                if "429" in error_message or "Too Many Requests" in error_message:
                    error_message = "Rate limit exceeded. Please wait a moment and try again."
                elif "401" in error_message or "Unauthorized" in error_message:
                    error_message = "Invalid API key. Please check your OpenAI API key."
                elif "400" in error_message:
                    error_message = "Invalid request. Please check your input."
                elif "500" in error_message:
                    error_message = "OpenAI service is temporarily unavailable. Please try again later."
                else:
                    error_message = f"Error generating response: {error_message}"
                
                yield f"data: {json.dumps({'error': error_message, 'type': 'error'})}\n\n"
                yield f"data: {json.dumps({'type': 'end'})}\n\n"

        response = StreamingHttpResponse(event_stream(), content_type='text/event-stream')
        response['Cache-Control'] = 'no-cache'
        response['Access-Control-Allow-Origin'] = '*'
        response['Access-Control-Allow-Headers'] = 'Cache-Control'
        return response

    def _create_documents_from_text(self):
        """
        Read the Whitepaper Tosi Growth Holding.docx file and store it in DocumentEmbedding as vector database.
        This method processes the DOCX file, splits it into chunks, and creates embeddings for each chunk.
        """
        try:
            # Path to the whitepaper file
            whitepaper_path = "Whitepaper Tosi Growth Holding.docx"
            
            if not os.path.exists(whitepaper_path):
                print(f"Error: Whitepaper file not found at {whitepaper_path}")
                return
            
            print(f"Reading whitepaper from: {whitepaper_path}")
            
            # Read the DOCX file
            doc = docx.Document(whitepaper_path)
            
            # Extract text from all paragraphs
            full_text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    full_text += paragraph.text.strip() + "\n\n"
            
            # Also extract text from tables if any
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text += cell.text.strip() + "\n"
                    full_text += "\n"
            
            if not full_text.strip():
                print("Warning: No text content found in the whitepaper file")
                return
            
            print(f"Extracted {len(full_text)} characters from whitepaper")
            
            # Create a Document object
            document = Document(
                page_content=full_text,
                metadata={
                    "source": "Whitepaper Tosi Growth Holding.docx",
                    "document_type": "whitepaper",
                    "company": "Tosi Growth Holding",
                    "extraction_method": "python-docx"
                }
            )
            
            # Split the document into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=300,
                chunk_overlap=100,
                length_function=len,
                separators=["\n\n", "\n", ". ", " ", ""]
            )
            
            documents = text_splitter.split_documents([document])
            print(f"Split document into {len(documents)} chunks")
            
            # Store documents in vector database
            create_document_embedding(documents)
            
            print("Successfully processed and stored whitepaper in vector database")
            
            # Also save the full content to JSON for backward compatibility
            self._save_whitepaper_to_json(full_text)
            
        except Exception as e:
            print(f"Error processing whitepaper: {e}")
            raise

    def _save_whitepaper_to_json(self, content):
        """
        Save the whitepaper content to JSON file for backward compatibility.
        """
        try:
            data = {
                "content": content,
                "source": "Whitepaper Tosi Growth Holding.docx",
                "processed_at": time.strftime("%Y-%m-%d %H:%M:%S")
            }
            
            with open('whitepaper_data.json', 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            
            print("Saved whitepaper content to whitepaper_data.json")
            
        except Exception as e:
            print(f"Error saving whitepaper to JSON: {e}")

    def search_whitepaper_embeddings(self, query: str, top_k: int = 5):
        """
        Search the whitepaper embeddings for relevant content based on a query.
        
        Args:
            query: The search query
            top_k: Number of top results to return
            
        Returns:
            List of relevant document chunks
        """
        try:
            from .models import get_pgvector_client, DocumentEmbedding
            
            # Check if we have any embeddings in the database
            if DocumentEmbedding.objects.count() == 0:
                print("No embeddings found in database. Using fallback content.")
                return []
            
            # Get the vector store client
            vector_store = get_pgvector_client()
            
            # Search for similar documents
            results = vector_store.similarity_search(query, k=top_k)

            print("================================================")
            print(results)
            print("================================================")
            
            return results
            
        except Exception as e:
            print(f"Error searching whitepaper embeddings: {e}")
            # Return empty list to fall back to full content
            return []
